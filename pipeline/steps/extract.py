"""
Ingest steps 1-2: file arrives -> extract raw text.
Supports: .pdf, .docx
"""
from pathlib import Path

import pdfplumber
from docx import Document


def extract(file_path: str) -> str:
    """Route to the right extractor based on file extension. Returns raw text."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(file_path)

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    elif ext == ".txt":
        return _extract_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext} (only .pdf, .docx, .txt)")


def _extract_pdf(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    doc = Document(path)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # include table content (docx tables aren't captured by .paragraphs)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python extract.py <file.pdf|file.docx>")
        sys.exit(1)

    raw = extract(sys.argv[1])
    print(f"--- Extracted {len(raw)} chars ---\n")
    print(raw[:2000])
